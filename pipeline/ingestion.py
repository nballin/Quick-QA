"""Load documents from disk or HuggingFace SQuAD."""

from __future__ import annotations

import hashlib
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


@dataclass
class Document:
    doc_id: str
    source: str
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


def _doc_id_from_path(path: Path) -> str:
    h = hashlib.sha256(str(path.resolve()).encode("utf-8")).hexdigest()
    return h[:8]


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_pdf(path: Path) -> str:
    from PyPDF2 import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text()
            if t:
                parts.append(t)
        except Exception as e:
            logger.warning("PDF page extract failed %s: %s", path, e)
    return "\n".join(parts)


def _read_docx(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())


def load_documents(corpus_dir: str | Path) -> list[Document]:
    """Recursively load .txt, .pdf, .docx from corpus_dir."""
    root = Path(corpus_dir).resolve()
    if not root.is_dir():
        logger.warning("Corpus directory does not exist: %s", root)
        return []

    supported = {".txt": _read_txt, ".pdf": _read_pdf, ".docx": _read_docx}
    docs: list[Document] = []

    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        ext = path.suffix.lower()
        if ext not in supported:
            logger.debug("Skipping unsupported extension: %s", path)
            continue
        try:
            text = supported[ext](path)
            if not text or not text.strip():
                logger.warning("Empty content, skipping: %s", path)
                continue
            try:
                src = str(path.relative_to(root))
            except ValueError:
                src = str(path)
            doc = Document(
                doc_id=_doc_id_from_path(path),
                source=src,
                text=text.strip(),
                metadata={
                    "filename": path.name,
                    "filetype": ext[1:],
                    "num_chars": len(text),
                },
            )
            docs.append(doc)
        except Exception as e:
            logger.warning("Failed to parse %s: %s — skipping", path, e)

    logger.info("Loaded %d documents from %s", len(docs), root)
    return docs


def load_squad_passages(split: str = "validation", n: int = 500) -> list[Document]:
    """Load unique contexts from rajpurkar/squad as Documents."""
    from datasets import load_dataset

    ds = load_dataset("rajpurkar/squad", split=split)
    seen: set[str] = set()
    docs: list[Document] = []
    for row in ds:
        ctx = row["context"].strip()
        if ctx in seen:
            continue
        seen.add(ctx)
        title = row.get("title", "")
        h = hashlib.sha256(ctx.encode("utf-8")).hexdigest()[:8]
        docs.append(
            Document(
                doc_id=h,
                source=f"squad::{title}",
                text=ctx,
                metadata={"title": title, "source": "squad"},
            )
        )
        if len(docs) >= n:
            break
    logger.info("Loaded %d unique SQuAD passages (split=%s)", len(docs), split)
    return docs
